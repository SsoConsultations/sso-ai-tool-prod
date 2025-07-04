import streamlit as st
import os
import io
import json
import re
import bcrypt
from datetime import datetime
import time
import base64
import subprocess # NEW: For calling external commands like unoconv
import tempfile # NEW: For creating temporary files

# --- Firebase Imports ---
import firebase_admin
from firebase_admin import credentials, auth, firestore
from firebase_admin import exceptions

# --- Google Drive Imports ---
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload # Added MediaIoBaseDownload for completeness

# --- AI & Document Processing Imports ---
from openai import OpenAI
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

# --- Streamlit Page Configuration (MUST BE THE FIRST ST COMMAND) ---
st.set_page_config(
    page_title="SSO Consultants AI Recruitment",
    page_icon="sso_logo.png",
    layout="wide"
)

# --- Custom CSS for Styling ---
st.markdown(
    """
    <style>
    /* Global base styling - pure white background, pure black text */
    body {
        background-color: #FFFFFF;
        color: #000000;
    }
    .stApp {
        background-color: #FFFFFF;
    }
    /* Main content area padding and width adjustment */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        padding-left: 5rem;
        padding-right: 5rem;
    }
    /* Specific styling for Streamlit widgets */
    .stButton>button {
        background-color: #4CAF50; /* Green button */
        color: white;
        border-radius: 5px;
        border: none;
        padding: 10px 20px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        margin: 4px 2px;
        cursor: pointer;
        -webkit-transition-duration: 0.4s; /* Safari */
        transition-duration: 0.4s;
        box-shadow: 0 4px 8px 0 rgba(0,0,0,0.2);
    }
    .stButton>button:hover {
        background-color: #45a049; /* Darker green on hover */
        box-shadow: 0 8px 16px 0 rgba(0,0,0,0.2);
    }
    .stTextInput>div>div>input {
        border-radius: 5px;
        border: 1px solid #ccc;
        padding: 10px;
    }
    .stFileUploader label {
        background-color: #f0f2f6; /* Light gray for file uploader */
        border-radius: 5px;
        padding: 10px;
        border: 1px dashed #ccc;
    }
    /* Styling for the navigation buttons */
    .stRadio > label {
        font-size: 1.1em;
        font-weight: bold;
        color: #4CAF50;
    }
    .stRadio div[role="radiogroup"] {
        display: flex;
        flex-direction: row;
        gap: 20px;
        justify-content: center; /* Center the radio buttons */
    }
    .stRadio div[role="radiogroup"] label {
        padding: 10px 15px;
        background-color: #e0ffe0; /* Very light green for tabs */
        border-radius: 10px;
        cursor: pointer;
        border: 1px solid #4CAF50;
        transition: all 0.3s ease-in-out;
        color: #4CAF50; /* Green text for tabs */
    }
    .stRadio div[role="radiogroup"] label:hover {
        background-color: #d0ffd0; /* Slightly darker green on hover */
    }
    .stRadio div[role="radiogroup"] label[data-baseweb="radio"] input:checked + span {
        background-color: #4CAF50 !important; /* Green for selected tab */
        color: white !important; /* White text for selected tab */
        border-color: #4CAF50 !important;
    }
    /* Make the radio button dot disappear */
    .stRadio [data-baseweb="radio"] > label::before {
        content: none !important;
    }
    .stRadio [data-baseweb="radio"] > label span {
        color: #4CAF50 !important; /* Ensure text remains green */
    }
    .stRadio [data-baseweb="radio"] > label[data-checked=true] span {
        color: white !important; /* White text for selected tab */
    }

    /* Style for the horizontal line divider */
    hr {
        border: 0;
        height: 1px;
        background-image: linear-gradient(to right, rgba(0, 0, 0, 0), rgba(0, 0, 0, 0.75), rgba(0, 0, 0, 0));
        margin: 2em 0;
    }

    /* Message styling */
    .st-emotion-cache-1ldf2b0 { /* Target for success messages */
        background-color: #e6ffe6; /* Light green */
        color: #1a661a; /* Dark green text */
        border: 1px solid #4CAF50;
        border-radius: 5px;
        padding: 10px;
    }
    .st-emotion-cache-1g8x1l5 { /* Target for error messages */
        background-color: #ffe6e6; /* Light red */
        color: #cc0000; /* Dark red text */
        border: 1px solid #ff0000;
        border-radius: 5px;
        padding: 10px;
    }

    /* Center images (adjust as needed) */
    .stImage {
        display: flex;
        justify-content: center;
        margin-top: 1rem;
        margin-bottom: 1rem;
    }

    /* For centering elements in columns */
    .st-emotion-cache-1nj6q9b { /* Targets Streamlit columns directly */
        justify-content: center;
        align-items: center;
        text-align: center;
    }

    /* Adjust padding for header */
    .css-1avcm0c { /* Header container */
        padding-top: 1rem;
        padding-bottom: 1rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- Initialize Firebase (Only once) ---
if not firebase_admin._apps:
    try:
        cred = credentials.Certificate('sso-consultants-firebase-adminsdk.json')
        firebase_admin.initialize_app(cred)
        db = firestore.client()
    except Exception as e:
        st.error(f"Error initializing Firebase: {e}. Please ensure 'sso-consultants-firebase-adminsdk.json' is correctly placed and accessible.")
        st.stop()

# --- Google Drive Service Setup ---
def get_drive_service():
    try:
        if "gcp_service_account" in st.secrets:
            info = st.secrets["gcp_service_account"]
            creds = service_account.Credentials.from_service_account_info(info)
        else:
            creds = service_account.Credentials.from_service_account_file(
                'gdrive_service_account.json',
                scopes=['https://www.googleapis.com/auth/drive']
            )
        return build('drive', 'v3', credentials=creds)
    except Exception as e:
        st.error(f"Error initializing Google Drive service: {e}")
        st.info("Please ensure your Google Drive service account key is correctly configured in `gdrive_service_account.json` or Streamlit secrets.")
        return None

# --- Session State Initialization ---
def init_session_state():
    if 'authenticated' not in st.session_state:
        st.session_state['authenticated'] = False
    if 'current_page' not in st.session_state:
        st.session_state['current_page'] = 'login'
    if 'username' not in st.session_state:
        st.session_state['username'] = None
    if 'is_admin' not in st.session_state:
        st.session_state['is_admin'] = False
    if 'current_admin_page' not in st.session_state:
        st.session_state['current_admin_page'] = 'generate'
    if 'google_drive_folder_id' not in st.session_state:
        st.session_state['google_drive_folder_id'] = os.environ.get('GOOGLE_DRIVE_FOLDER_ID', 'YOUR_GOOGLE_DRIVE_FOLDER_ID_HERE')
        if st.session_state['google_drive_folder_id'] == 'YOUR_GOOGLE_DRIVE_FOLDER_ID_HERE':
            st.warning("Please set your Google Drive Folder ID in environment variables or directly in the code for saving reports.")

init_session_state()

# --- User Management Functions ---
def register_user(email, password):
    try:
        user = auth.create_user(email=email, password=password)
        db.collection('users').document(user.uid).set({
            'email': email,
            'role': 'user',
            'created_at': firestore.SERVER_TIMESTAMP
        })
        return user.uid
    except Exception as e:
        st.error(f"Error registering user: {e}")
        return None

def login_user(email, password):
    try:
        user_record = auth.get_user_by_email(email)
        user_doc = db.collection('users').document(user_record.uid).get()
        if user_doc.exists:
            user_data = user_doc.to_dict()
            st.session_state['authenticated'] = True
            st.session_state['username'] = email
            st.session_state['is_admin'] = (user_data.get('role') == 'admin')
            return True
        else:
            st.error("User not found in database.")
            return False
    except firebase_admin.exceptions.FirebaseError as e:
        if "EMAIL_NOT_FOUND" in str(e) or "INVALID_PASSWORD" in str(e):
            st.error("Invalid email or password.")
        else:
            st.error(f"Login error: {e}")
        return False
    except Exception as e:
        st.error(f"An unexpected error occurred during login: {e}")
        return False

def logout_user():
    st.session_state['authenticated'] = False
    st.session_state['username'] = None
    st.session_state['is_admin'] = False
    st.session_state['current_page'] = 'login'
    st.session_state['current_admin_page'] = 'generate'
    st.rerun()

def get_all_users():
    users_ref = db.collection('users')
    docs = users_ref.stream()
    users = []
    for doc in docs:
        user_data = doc.to_dict()
        user_data['uid'] = doc.id
        users.append(user_data)
    return users

def update_user_role(uid, new_role):
    try:
        db.collection('users').document(uid).update({'role': new_role})
        st.success(f"User {uid} role updated to {new_role}")
    except Exception as e:
        st.error(f"Error updating user role: {e}")

def delete_user(uid):
    try:
        auth.delete_user(uid)
        db.collection('users').document(uid).delete()
        st.success(f"User {uid} deleted successfully.")
    except Exception as e:
        st.error(f"Error deleting user: {e}")

# --- Google Drive Functions ---
def upload_file_to_drive(file_content_io, file_name, mime_type, folder_id):
    drive_service = get_drive_service()
    if not drive_service:
        return None

    file_metadata = {
        'name': file_name,
        'parents': [folder_id],
        'mimeType': mime_type
    }
    media = MediaIoBaseUpload(file_content_io, mimetype=mime_type, resumable=True)
    try:
        file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        st.success(f"Report '{file_name}' uploaded to Google Drive!")
        return file.get('id')
    except Exception as e:
        st.error(f"Failed to upload '{file_name}' to Google Drive: {e}")
        return None

def list_files_in_drive_folder(folder_id):
    drive_service = get_drive_service()
    if not drive_service:
        return []

    try:
        results = drive_service.files().list(
            q=f"'{folder_id}' in parents and trashed = false",
            fields="nextPageToken, files(id, name, mimeType, modifiedTime)").execute()
        items = results.get('files', [])
        return items
    except Exception as e:
        st.error(f"Error listing files from Google Drive: {e}")
        return []

def download_file_from_drive(file_id, file_name):
    drive_service = get_drive_service()
    if not drive_service:
        return None

    try:
        request = drive_service.files().get_media(fileId=file_id)
        file_content = io.BytesIO()
        downloader = MediaIoBaseDownload(file_content, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        file_content.seek(0)
        return file_content
    except Exception as e:
        st.error(f"Error downloading file '{file_name}' from Google Drive: {e}")
        return None

# --- Document Processing Functions ---

def add_styled_paragraph(document, text, style='Normal', font_size=12, bold=False, italic=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph(text, style=style)
    run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.alignment = alignment
    return paragraph

def add_styled_heading(document, text, level=1, font_size=16, color='000000', alignment=WD_ALIGN_PARAGRAPH.LEFT):
    heading = document.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.size = Pt(font_size)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    heading.alignment = alignment
    return heading

def extract_text_from_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

# --- OpenAI API Interaction ---
def get_openai_response(prompt_template, jd_text, cv_texts, model="gpt-4o", max_tokens=1500, temperature=0.7):
    # Ensure OPENAI_API_KEY is set in environment variables or st.secrets
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY") or st.secrets["openai_api_key"])
    except KeyError:
        st.error("OpenAI API key not found. Please set OPENAI_API_KEY in your environment variables or Streamlit secrets.")
        return "Error: OpenAI API key not configured."
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {e}")
        return "Error: OpenAI client initialization failed."

    try:
        messages = [
            {"role": "system", "content": "You are a highly skilled AI specializing in Job Description and CV analysis. Provide accurate, concise, and professional assessments."},
            {"role": "user", "content": prompt_template.format(jd_text=jd_text, cv_texts=cv_texts)}
        ]
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error getting response from OpenAI API: {e}")
        st.toast(f"AI API Error: {e}", icon="🚨")
        return f"Error: Could not get response from AI. Details: {e}"

# --- Report Generation Prompts ---
JOB_DESCRIPTION_ANALYSIS_PROMPT = """
Analyze the following Job Description (JD) and provide a concise summary of its key requirements, responsibilities, and preferred qualifications.
JD: {jd_text}
"""

OVERALL_CV_ANALYSIS_PROMPT = """
Based on the provided Job Description (JD) and a collection of anonymized CVs, give an overall summary of the collective strengths and weaknesses of the candidates relative to the JD. Highlight common trends, gaps, or outstanding qualities observed across the CVs.

JD: {jd_text}
CVs: {cv_texts}
"""

INDIVIDUAL_CV_COMPARISON_PROMPT = """
Compare the following CV against the provided Job Description (JD). For this specific CV, provide:
1. Key strengths/alignments with the JD.
2. Key areas of improvement/gaps in the CV relative to the JD.
3. An overall fit score (e.g., 1-10 or Poor/Fair/Good/Excellent) and a brief justification.

JD: {jd_text}
CV: {cv_texts}
"""

# --- Main Pages ---

def login_page():
    st.title("Login to SSO Consultants AI Recruitment")
    st.subheader("Welcome Back!")
    email = st.text_input("Email", key="login_email")
    password = st.text_input("Password", type="password", key="login_password")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Login", key="login_button_main"):
            if login_user(email, password):
                st.success("Logged in successfully!")
                st.session_state['current_page'] = 'dashboard'
                st.rerun()
            else:
                st.error("Invalid credentials.")
    with col2:
        if st.button("Register", key="register_button_main"):
            st.session_state['current_page'] = 'register'
            st.rerun()

def register_page():
    st.title("Register for SSO Consultants AI Recruitment")
    st.subheader("Create Your Account")
    email = st.text_input("Email", key="register_email")
    password = st.text_input("Password", type="password", key="register_password")
    confirm_password = st.text_input("Confirm Password", type="password", key="confirm_password")

    if st.button("Submit Registration", key="submit_registration"):
        if password == confirm_password:
            if register_user(email, password):
                st.success("Registration successful! You can now login.")
                st.session_state['current_page'] = 'login'
                st.rerun()
            else:
                st.error("Registration failed. Please try again.")
        else:
            st.error("Passwords do not match.")

    if st.button("Back to Login", key="back_to_login"):
        st.session_state['current_page'] = 'login'
        st.rerun()

def generate_comparative_report_page():
    st.header("Generate Comparative Report")

    # Upload Job Description
    jd_file = st.file_uploader("Upload Job Description (PDF only)", type=["pdf"], key="jd_uploader")
    if jd_file:
        st.info("JD file uploaded.")
        st.session_state['jd_uploaded'] = True
        st.session_state['jd_name'] = jd_file.name
        st.session_state['jd_text'] = extract_text_from_pdf(jd_file)

    # Upload CVs
    cv_files = st.file_uploader("Upload CVs (PDF only, multiple allowed)", type=["pdf"], accept_multiple_files=True, key="cv_uploader")
    if cv_files:
        st.info(f"{len(cv_files)} CV(s) uploaded.")
        st.session_state['cv_uploaded'] = True
        st.session_state['cv_data'] = {}
        for cv_file in cv_files:
            st.session_state['cv_data'][cv_file.name] = extract_text_from_pdf(cv_file)

    st.markdown("---")

    if st.button("Generate Report", key="generate_report_button"):
        if 'jd_text' not in st.session_state or not st.session_state['jd_text']:
            st.error("Please upload a Job Description PDF first.")
            return
        if 'cv_data' not in st.session_state or not st.session_state['cv_data']:
            st.error("Please upload at least one CV PDF first.")
            return

        with st.spinner("Generating report... This may take a few moments."):
            jd_text = st.session_state['jd_text']
            cv_data = st.session_state['cv_data']
            cv_names = list(cv_data.keys())

            # Create a new Word document
            document = Document()
            document.add_heading('JD-CV Comparative Analysis Report', level=0)
            add_styled_paragraph(document, f"Generated by: {st.session_state['username']}", font_size=10, color='666666')
            add_styled_paragraph(document, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", font_size=10, color='666666')
            add_styled_paragraph(document, f"Job Description File: {st.session_state['jd_name']}", font_size=10, color='666666')
            add_styled_paragraph(document, f"CV Files Analyzed: {', '.join(cv_names)}", font_size=10, color='666666')
            document.add_page_break()

            # --- Job Description Analysis ---
            add_styled_heading(document, "Job Description Analysis", level=1, color='2F5496')
            jd_analysis_response = get_openai_response(JOB_DESCRIPTION_ANALYSIS_PROMPT, jd_text, "")
            add_styled_paragraph(document, jd_analysis_response)
            document.add_page_break()

            # --- Overall CV Analysis ---
            add_styled_heading(document, "Overall CV Analysis (Summary)", level=1, color='2F5496')
            all_cvs_text = "\n\n".join([f"CV for {name}:\n{text}" for name, text in cv_data.items()])
            overall_cv_response = get_openai_response(OVERALL_CV_ANALYSIS_PROMPT, jd_text, all_cvs_text)
            add_styled_paragraph(document, overall_cv_response)
            document.add_page_break()

            # --- Individual CV Comparison ---
            add_styled_heading(document, "Individual CV Comparison", level=1, color='2F5496')
            for cv_name, cv_text in cv_data.items():
                add_styled_heading(document, f"{cv_name} Comparison", level=2, color='1F4E79')
                individual_comparison_response = get_openai_response(INDIVIDUAL_CV_COMPARISON_PROMPT, jd_text, cv_text)
                add_styled_paragraph(document, individual_comparison_response)
                document.add_paragraph("")
            document.add_page_break()

            # --- Save DOCX to a BytesIO object first ---
            docx_buffer = io.BytesIO()
            document.save(docx_buffer)
            docx_buffer.seek(0)

            report_filename_base = f"JD_CV_Analysis_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            docx_filename = f"{report_filename_base}.docx"
            pdf_filename = f"{report_filename_base}.pdf"

            # --- Upload DOCX to Google Drive ---
            if st.session_state['google_drive_folder_id'] and st.session_state['google_drive_folder_id'] != 'YOUR_GOOGLE_DRIVE_FOLDER_ID_HERE':
                docx_buffer_for_upload = io.BytesIO(docx_buffer.getvalue())
                upload_file_to_drive(docx_buffer_for_upload, docx_filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', st.session_state['google_drive_folder_id'])
            else:
                st.warning("Google Drive folder ID not configured. Report will not be saved to Drive.")

            # --- Convert DOCX to PDF for Preview ---
            pdf_buffer = None
            with tempfile.TemporaryDirectory() as tmpdir:
                temp_docx_path = os.path.join(tmpdir, docx_filename)
                temp_pdf_path = os.path.join(tmpdir, pdf_filename)

                with open(temp_docx_path, "wb") as f:
                    f.write(docx_buffer.getvalue())

                try:
                    subprocess.run(
                        ["unoconv", "-f", "pdf", "-o", temp_pdf_path, temp_docx_path],
                        check=True,
                        capture_output=True,
                        text=True
                    )
                    st.success("DOCX converted to PDF successfully for preview!")

                    with open(temp_pdf_path, "rb") as f:
                        pdf_buffer = io.BytesIO(f.read())
                        pdf_buffer.seek(0)

                except FileNotFoundError:
                    st.error("`unoconv` or LibreOffice not found. PDF preview will not be available. Please install LibreOffice and unoconv on your system.")
                    st.info("Example install: `sudo apt install libreoffice unoconv` (Ubuntu/Debian)")
                except subprocess.CalledProcessError as e:
                    st.error(f"Error during DOCX to PDF conversion: {e.stderr}")
                    st.toast("PDF conversion failed!", icon="❌")
                except Exception as e:
                    st.error(f"An unexpected error occurred during PDF conversion: {e}")
                    st.toast("PDF conversion failed!", icon="❌")

            st.success("Report generation complete!")

            # --- Display PDF Preview (if available) ---
            if pdf_buffer:
                st.subheader("Report Preview")
                st.download_button(
                    label="View Report Preview (PDF)",
                    data=pdf_buffer,
                    file_name=pdf_filename,
                    mime="application/pdf",
                    key="pdf_preview_button",
                    help="Click to view the report as a PDF in your browser."
                )
                st.info("The PDF preview button above will typically open the report in your browser's PDF viewer. "
                        "Depending on your browser, it might display directly here or open in a new tab.")
            else:
                st.warning("PDF preview could not be generated. Please check server logs for `unoconv` issues.")

            # --- Provide DOCX Download Option ---
            st.subheader("Download Original Report")
            st.download_button(
                label="Download Full Report (DOCX)",
                data=docx_buffer,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx_button",
                help="Click to download the report in Microsoft Word (.docx) format."
            )
            st.success("Report ready for download!")
            st.info("The report has also been uploaded to Google Drive (if configured).")

def show_all_reports_page():
    st.header("All Generated Reports")
    folder_id = st.session_state['google_drive_folder_id']

    if folder_id == 'YOUR_GOOGLE_DRIVE_FOLDER_ID_HERE':
        st.warning("Google Drive folder ID not configured. Cannot list reports from Drive.")
        return

    reports = list_files_in_drive_folder(folder_id)

    if not reports:
        st.info("No reports found in Google Drive.")
        return

    report_data = []
    for item in reports:
        report_data.append({
            "File Name": item['name'],
            "Modified Time": datetime.fromisoformat(item['modifiedTime'].replace('Z', '+00:00')).strftime('%Y-%m-%d %H:%M:%S'),
            "File ID": item['id'],
            "Mime Type": item['mimeType']
        })

    df_reports = pd.DataFrame(report_data)
    df_reports = df_reports.sort_values(by="Modified Time", ascending=False).reset_index(drop=True)

    st.dataframe(df_reports, use_container_width=True)

    st.subheader("Download a Report")
    selected_report_name = st.selectbox("Select a report to download", df_reports['File Name'].tolist(), key="select_report_download")

    if selected_report_name:
        selected_report = df_reports[df_reports['File Name'] == selected_report_name].iloc[0]
        file_id = selected_report['File ID']
        file_name = selected_report['File Name']
        mime_type = selected_report['Mime Type']

        if st.button(f"Download '{file_name}'", key="download_selected_report_button"):
            with st.spinner(f"Downloading '{file_name}'..."):
                downloaded_file_io = download_file_from_drive(file_id, file_name)
                if downloaded_file_io:
                    st.download_button(
                        label="Click to Download",
                        data=downloaded_file_io,
                        file_name=file_name,
                        mime=mime_type,
                        key=f"download_button_{file_id}"
                    )
                    st.success(f"'{file_name}' ready for download!")
                else:
                    st.error("Failed to download file.")

def manage_users_page():
    st.header("Manage Users")
    if not st.session_state['is_admin']:
        st.warning("You do not have administrative privileges to manage users.")
        return

    users = get_all_users()
    if not users:
        st.info("No users registered yet.")
        return

    user_data = []
    for user in users:
        user_data.append({
            "UID": user['uid'],
            "Email": user['email'],
            "Role": user.get('role', 'user')
        })

    df_users = pd.DataFrame(user_data)
    st.dataframe(df_users, use_container_width=True)

    st.subheader("Update User Role")
    user_email_to_update = st.selectbox("Select user by Email to update role", df_users['Email'].tolist(), key="update_role_email_select")
    new_role = st.selectbox("New Role", ["user", "admin"], key="new_role_select")

    if st.button("Update Role", key="update_role_button"):
        selected_uid = df_users[df_users['Email'] == user_email_to_update]['UID'].iloc[0]
        update_user_role(selected_uid, new_role)
        st.rerun()

    st.subheader("Delete User")
    user_email_to_delete = st.selectbox("Select user by Email to delete", df_users['Email'].tolist(), key="delete_user_email_select")

    if st.button("Delete User", key="delete_user_button", help="This action is irreversible!"):
        # Added a confirmation step
        st.warning(f"Are you sure you want to delete user {user_email_to_delete}? This action is irreversible.")
        col_yes, col_no = st.columns(2)
        with col_yes:
            if st.button("Yes, Delete", key="confirm_delete_button_yes"):
                selected_uid = df_users[df_users['Email'] == user_email_to_delete]['UID'].iloc[0]
                delete_user(selected_uid)
                st.rerun()
        with col_no:
            if st.button("No, Cancel", key="confirm_delete_button_no"):
                st.info("User deletion cancelled.")


# --- Main App Logic ---
def main():
    # Logo and App Title - Always visible regardless of login state
    col_logo, col_title = st.columns([1, 4])
    with col_logo:
        st.image("sso_logo.png", width=100, use_container_width=False)
    with col_title:
        st.title("AI Recruitment Platform")
        st.caption("Powered by SSO Consultants")

    if st.session_state['authenticated']:
        st.sidebar.header(f"Welcome, {st.session_state['username']}!")
        if st.session_state['is_admin']:
            st.sidebar.markdown("**Admin Panel**")
            admin_page_selection = st.sidebar.radio(
                "Navigation",
                ('Generate Report', 'All Reports', 'Manage Users'),
                key="admin_nav_radio",
                index=0 if st.session_state['current_admin_page'] == 'generate' else (1 if st.session_state['current_admin_page'] == 'reports' else 2)
            )
            if admin_page_selection == 'Generate Report':
                st.session_state['current_admin_page'] = 'generate'
            elif admin_page_selection == 'All Reports':
                st.session_state['current_admin_page'] = 'reports'
            elif admin_page_selection == 'Manage Users':
                st.session_state['current_admin_page'] = 'manage_users'
        else:
            st.sidebar.markdown("**User Panel**")
            user_page_selection = st.sidebar.radio(
                "Navigation",
                ('Generate Report', 'All Reports'),
                key="user_nav_radio",
                index=0 if st.session_state['current_admin_page'] == 'generate' else 1
            )
            if user_page_selection == 'Generate Report':
                st.session_state['current_admin_page'] = 'generate'
            elif user_page_selection == 'All Reports':
                st.session_state['current_admin_page'] = 'reports'

        st.sidebar.write("---")
        if st.sidebar.button("Logout", key="logout_button"):
            logout_user()

        if st.session_state['is_admin']:
            if st.session_state['current_admin_page'] == 'generate':
                generate_comparative_report_page()
            elif st.session_state['current_admin_page'] == 'reports':
                show_all_reports_page()
            elif st.session_state['current_admin_page'] == 'manage_users':
                manage_users_page()
        else:
            if st.session_state['current_admin_page'] == 'generate':
                generate_comparative_report_page()
            elif st.session_state['current_admin_page'] == 'reports':
                show_all_reports_page()

    else:
        if st.session_state['current_page'] == 'login':
            login_page()
        elif st.session_state['current_page'] == 'register':
            register_page()

st.markdown(
    """
    <div style="
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        text-align: center;
        color: #4CAF50 !important;
        padding: 10px;
        background-color: #FFFFFF;
        font-size: 0.8em;
        border-top: 1px solid #eee;
        z-index: 9999;
    ">
        <p>&copy; 2024 SSO Consultants. All rights reserved. Made for AI Recruitment.</p>
    </div>
    """,
    unsafe_allow_html=True
)

if __name__ == "__main__":
    main()
